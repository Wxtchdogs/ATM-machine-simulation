from docx import Document

#Create new docx file
document = Document()

#Contents of the page
document.add_heading('This is my docx file for assessment 3', level = 0)
document.add_heading('This is the subheading', level = 1)
document.add_paragraph('This is my first paragraph')
document.add_picture('cat.jpeg')

#Save the docx file
document.save('mydoc.docx')

#Obtain contents of the docx file function
def obtainText(docFile):
    document = Document(docFile)
    contents = []
    for item in document.paragraphs:
        contents.append(item.text)
    #Add each item to new line
    return '\n'.join(contents)

#Print the file's contents
print(obtainText('mydoc.docx'))
